from .precision_pdf_model import PrecisionPDFModel
from .miner_pdf_model import MinerPDFModel
from .ms_model import MSWordModel
from typing import Dict, List
import docx
from docx.enum.style import WD_STYLE_TYPE


class PrecisionPDFToMSWord:
    def convert(self, input_model: PrecisionPDFModel, output_model: MSWordModel)-> None:
        pages: List[Dict] = input_model.pdf_json["pages"]
        output_model.doc = docx.Document()
        styles = output_model.doc.styles
        self.LABELS = {
            "text": "Normal",
            "header": "Heading 1",
            "list": "List",
            "table": "Normal",
            # "figure": "Normal",
        }
        for page in pages:
            self.add_page(output_model.doc, page, styles)
      

    def add_page(self, doc: docx.Document, page: Dict, styles)-> None:
        for regions in page["regions"]:
            label = regions["label"]
            text = regions["text"]
            if label != "figure":
                doc.add_paragraph(text).style = styles[self.LABELS[label]]
        doc.add_page_break()

class MinerPDFToMSWord:
    def convert(self, input_model: MinerPDFModel, output_model: MSWordModel)-> None:
        pages: List[Dict] = input_model.pdf_json["pages"]
        output_model.doc = docx.Document()
        styles = output_model.doc.styles
        self.LABELS = {
            "text": "Normal",
            "header": "Heading 1",
            "list": "List",
            "table": "Normal",
            # "figure": "Normal",
        }
        for page in pages:
            self.add_page(output_model.doc, page, styles)
      

    def add_page(self, doc: docx.Document, page: Dict, styles)-> None:
        for regions in page["regions"]:
            label = regions["label"]
            text = regions["text"]
            if label != "figure":
                doc.add_paragraph(text).style = styles[self.LABELS[label]]
        doc.add_page_break()