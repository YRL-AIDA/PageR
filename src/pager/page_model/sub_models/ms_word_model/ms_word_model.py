from ..base_sub_model import BaseSubModel, BaseExtractor, BaseConverter
from typing import Dict, List
import docx
from ..dtype import Block
from docx.enum.style import WD_STYLE_TYPE


class MSWordModel(BaseSubModel):
    def __init__(self) -> None:
        super().__init__()
        self.doc: docx.Document
        
    def from_dict(self, input_model_dict: Dict):
        self.doc = docx.Document()
        styles = dict()
        for st in input_model_dict['styles']:
            styles[st['id']] = {"name": st['name'], "font": st['font']}
        
        for block in input_model_dict["blocks"]:
            self.doc.add_paragraph(block["text"]).style = styles[block["style"]]["name"]

        for key, st in styles.items():
            self.doc.styles[key].font.name = st["font"]["name"]
            self.doc.styles[key].font.size = st["font"]["size"]
            self.doc.styles[key].font.bold = st["font"]["bold"]
            self.doc.styles[key].font.italic = st["font"]["italic"]
            self.doc.styles[key].font.underline = st["font"]["underline"]
    
    def save_doc(self, path):
        self.doc.save(path)

    def to_dict(self) -> Dict:
        styles_paragraphs = []
        for pr in self.doc.paragraphs:
            if not pr.style in styles_paragraphs:
                styles_paragraphs.append(pr.style)
      
        return {"styles": [self.style_to_dict(st) for st in self.doc.styles if st in styles_paragraphs],
                "blocks": [self.paragraph_to_dict(pr) for pr in self.doc.paragraphs]}

    def read_from_file(self, path_file: str) -> None:
        self.doc = docx.Document(path_file)

    def clean_model(self):
        self.doc = None

    def style_to_dict(self, st) -> Dict:
        return {
            "id": st.style_id,
            "name": st.name,
            "font": {
                "name": st.font.name,
                "size": st.font.size,
                "bold": st.font.bold,
                "italic": st.font.italic,
                "underline": st.font.underline}
            }
    
    def paragraph_to_dict(self, pr) -> Dict:
        return {
            "style": pr.style.style_id,
            "text": pr.text,
            # "runs": [self.run_to_dict(r) for r in pr.runs]
            }
            
    def run_to_dict(self, r) -> Dict:
        return {
            "style_id": r.style.style_id,
            "text": r.text,
            "font": {
                "name": r.font.name,
                "size": r.font.size,
                "bold": r.font.bold,
                "italic": r.font.italic,
                "underline": r.font.underline}
            }
    
    
class PhisicalToMSWord(BaseConverter):
    def convert(self, input_model: BaseSubModel, output_model: BaseSubModel)-> None:
        blocks: List[Block] = input_model.blocks
        output_model.doc = docx.Document()
        styles = output_model.doc.styles
        LABELS = {
            "text": "Normal",
            "header": "Heading 1",
            "list": "List",
            "table": "Normal",
            # "figure": "Normal",
        }
        for block in blocks:
            if block.label != "figure":
                output_model.doc.add_paragraph(block.get_text()).style = styles[LABELS[block.label]]
    



